# LangChain compatibility fix - must be imported first
import langchain_compat  # noqa: F401

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader
from langchain_core.documents import Document
from docx import Document as DocxDocument
from typing import List, Optional
from pathlib import Path
import logging

from vectorstore import vector_store

logger = logging.getLogger(__name__)


def load_pdf(file_path: str) -> List[Document]:
    """Load PDF file and return Document objects"""
    try:
        loader = PyPDFLoader(file_path)
        documents = loader.load()
        logger.info(f"Loaded {len(documents)} pages from PDF: {file_path}")
        return documents
    except Exception as e:
        logger.error(f"Error loading PDF {file_path}: {e}")
        raise


def load_docx(file_path: str) -> List[Document]:
    """Load DOCX file and return Document objects"""
    try:
        doc = DocxDocument(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
        
        document = Document(
            page_content=text,
            metadata={"source": file_path}
        )
        logger.info(f"Loaded DOCX: {file_path}")
        return [document]
    except Exception as e:
        logger.error(f"Error loading DOCX {file_path}: {e}")
        raise


def load_txt(file_path: str) -> List[Document]:
    """Load TXT file and return Document objects"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        document = Document(
            page_content=text,
            metadata={"source": file_path}
        )
        logger.info(f"Loaded TXT: {file_path}")
        return [document]
    except Exception as e:
        logger.error(f"Error loading TXT {file_path}: {e}")
        raise


def load_document_by_type(file_path: str) -> List[Document]:
    """Load document based on file extension"""
    path = Path(file_path)
    extension = path.suffix.lower()
    
    loaders = {
        '.pdf': load_pdf,
        '.docx': load_docx,
        '.doc': load_docx,
        '.txt': load_txt
    }
    
    loader = loaders.get(extension)
    if not loader:
        raise ValueError(f"Unsupported file type: {extension}")
    
    return loader(file_path)


def load_and_chunk_document(
    file_path: str,
    doc_id: str,
    case_id: Optional[str] = None,
    chunk_size: int = 1200,
    chunk_overlap: int = 200
) -> List[Document]:
    """
    Load a document, split it into chunks, and add metadata
    
    Args:
        file_path: Path to the document file
        doc_id: Unique document identifier
        case_id: Optional case identifier
        chunk_size: Size of text chunks
        chunk_overlap: Overlap between chunks
    
    Returns:
        List of chunked Document objects with metadata
    """
    try:
        # Load document
        documents = load_document_by_type(file_path)
        
        # Initialize text splitter
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        
        # Split documents into chunks
        chunks = text_splitter.split_documents(documents)
        
        # Add metadata to each chunk
        for i, chunk in enumerate(chunks):
            chunk.metadata.update({
                "doc_id": doc_id,
                "chunk_index": i,
                "source": file_path
            })
            if case_id:
                chunk.metadata["case_id"] = case_id
        
        logger.info(f"Created {len(chunks)} chunks for doc_id: {doc_id}")
        return chunks
        
    except Exception as e:
        logger.error(f"Error processing document {file_path}: {e}")
        raise


def index_document(
    file_path: str,
    doc_id: str,
    case_id: Optional[str] = None
) -> bool:
    """
    Load, chunk, and index a document into FAISS
    
    Args:
        file_path: Path to the document file
        doc_id: Unique document identifier
        case_id: Optional case identifier
    
    Returns:
        True if successful, False otherwise
    """
    try:
        # Load and chunk document
        chunks = load_and_chunk_document(file_path, doc_id, case_id)
        
        if not chunks:
            logger.warning(f"No chunks created for doc_id: {doc_id}")
            return False
        
        # Create FAISS vectorstore
        vector_store.create_vectorstore(chunks, doc_id)
        
        logger.info(f"Successfully indexed doc_id: {doc_id} with {len(chunks)} chunks")
        return True
        
    except Exception as e:
        logger.error(f"Error indexing document {doc_id}: {e}")
        return False